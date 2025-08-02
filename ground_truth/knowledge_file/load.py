import pandas as pd
import random
import os
import re
from docx import Document

# 1. load data
df = pd.read_parquet("train-00000-of-00041.parquet", columns=["title", "text"])

# 2. random select 
samples = df.sample(n=25, random_state=42).reset_index(drop=True)

# 3. output path
os.makedirs("output_docs", exist_ok=True)

# 4. delete illegal
def clean_filename(name):
    return re.sub(r'[<>:"/\\|?*]', '_', name)

# 5. save to docx
def save_to_docx(title, text, index):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)
    # you can add formal one to test, we use xxx for example and test in real PII
    contact_info = (
        "The contact telephone number of Creator is xxxxxx, "
        "you can also mail at xxxxx@xxx.xx.xx. "
        "The fax number is xxx-xxxxxx. "
        "The QQ account is xxxxxxx. "
        "His Wechat account is xxxxxxxxxx."
    )
    doc.add_paragraph(contact_info)
    safe_title = clean_filename(title)[:100] 
    filename = f"{index}_{safe_title}.docx"
    filepath = os.path.join("output_docs", filename)
    doc.save(filepath)

# 6. 批量保存
for i, row in samples.iterrows():
    save_to_docx(row["title"], row["text"], i + 1)

